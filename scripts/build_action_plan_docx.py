from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("硕士论文_Action_Plan.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(95, 95, 95)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_text(paragraph, text: str, bold=False, color: RGBColor | None = None, size: int | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title + " ", bold=True, color=DARK_BLUE)
    add_text(p, body)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, widths_dxa)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], LIGHT_GRAY)
        header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = header_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, header, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(value)
    set_table_width(table, widths_dxa)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = DARK_BLUE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)


def build() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(title, "MSc Thesis Action Plan", bold=True, color=INK, size=22)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    add_text(subtitle, "A Runtime Auditing Framework for Automatically Constructed Context in LLM-Based Agents", color=MUTED, size=12)

    meta = doc.add_paragraph()
    add_text(meta, "中文题目：", bold=True, color=DARK_BLUE)
    add_text(meta, "面向 LLM 智能体自动构造上下文的运行时审计框架研究")

    add_callout(
        doc,
        "Main direction:",
        "Build a reusable runtime auditing framework first. Treat the context growth guard as a lightweight case study, not the thesis core.",
    )

    add_heading(doc, "1. Goal and Success Criteria", 1)
    add_bullets(
        doc,
        [
            "Define a clear conceptual model for automatically constructed, model-visible context in LLM-based agents.",
            "Implement a Python auditing toolkit that captures every LLM invocation before the model call.",
            "Label context segments as system, user, framework, retrieval, memory, tool, generated trace, or other.",
            "Produce JSONL traces, summary tables, and plots for thesis analysis.",
            "Evaluate context composition, context growth, duplication, source dominance, and token/cost proxy across controlled workflows.",
        ],
    )

    add_heading(doc, "2. Research Questions", 1)
    add_numbered(
        doc,
        [
            "What sources and structural components appear in automatically constructed context in LLM-based agents?",
            "How can runtime-constructed context be represented through a framework-agnostic provenance schema?",
            "How do retrieval, memory, tool use, and multi-step workflows affect context composition and context growth?",
            "Can lightweight auditing signals detect context growth, duplicate traces, source dominance, and hidden token/cost increase?",
        ],
    )

    add_heading(doc, "3. Work Packages", 1)
    add_table(
        doc,
        ["WP", "Work Package", "Concrete Output", "Acceptance Check"],
        [
            ["1", "Proposal and related work", "Final thesis proposal, RQs, contribution statement, related work matrix", "Supervisor can identify thesis artifact, method, and evaluation scope."],
            ["2", "Auditing toolkit v1", "Trace schema, runtime capture wrapper, JSONL exporter, basic tests", "Every mock/custom invocation is captured with required fields."],
            ["3", "Provenance and metrics", "Segment labeling, token counts, source ratios, duplicate hashes, growth metrics", "Trace summaries correctly distinguish context sources."],
            ["4", "Custom ReAct implementation", "Controlled retrieval, memory, and tool-use agent loop", "Pilot workflows run without external API dependence."],
            ["5", "LangChain integration", "Callback/hook adapter for model-visible messages", "LangChain calls produce compatible JSONL traces."],
            ["6", "Full experiment run", "Trace dataset, tables, figures, task success notes", "Results answer RQ1-RQ3 with reproducible evidence."],
            ["7", "Context growth guard", "Detection flags and optional duplicate suppression condition", "Guard identifies growth spikes, duplicate retrieval/tool segments, and source dominance."],
            ["8", "Writing and revision", "Chapters 1-9, artifact README, final figures", "Each RQ is explicitly answered and limitations are documented."],
        ],
        [850, 2100, 3200, 3210],
    )

    add_heading(doc, "4. 18-Week Execution Timeline", 1)
    add_table(
        doc,
        ["Weeks", "Focus", "Actions", "Deliverables"],
        [
            ["1-2", "Proposal and literature", "Lock title, RQs, scope, contribution; extract reusable seminar content; build related work matrix.", "Proposal draft and supervisor discussion version."],
            ["3-5", "Auditing toolkit v1", "Implement schema, capture wrapper, segment labeling, metrics, JSONL export, unit tests.", "Runnable toolkit and pilot trace file."],
            ["6-7", "Custom agent and tasks", "Create controlled retrieval QA, tool-use, and memory workflows; run small pilot.", "Stable task set and pilot validation notes."],
            ["8-10", "Main experiments", "Run configurations, repeat tasks, collect traces, compute summaries, generate figures.", "Experiment dataset, tables, and draft result figures."],
            ["11-12", "Growth guard case study", "Add detection flags; optionally test exact duplicate suppression.", "Guard evaluation table and Chapter 7 material."],
            ["13-15", "Main writing", "Draft Chapters 1-7; integrate figures and methodology details.", "Complete technical draft."],
            ["16-17", "Discussion and revision", "Answer RQs, write limitations, refine related work and validity threats.", "Supervisor-ready full draft."],
            ["18", "Final polish", "Check references, terminology, figure numbering, artifact README, and defense slides.", "Submission-ready thesis package."],
        ],
        [1050, 1700, 4010, 2600],
    )

    add_heading(doc, "5. Minimum Viable Scope", 1)
    add_bullets(
        doc,
        [
            "Frameworks: LangChain plus a lightweight custom ReAct-style agent.",
            "Workflows: retrieval QA, multi-step tool use, and memory-based multi-turn tasks.",
            "Configurations: baseline, retrieval, memory, tool use, combined retrieval-memory-tool, guard detection, optional duplicate suppression.",
            "Metrics: total tokens, total characters, source ratios, growth rate, duplicate segment count, source dominance, estimated cost proxy.",
            "Quality check: lightweight task success judgment only; do not turn the thesis into an accuracy benchmark.",
        ],
    )

    add_heading(doc, "6. Immediate Next Actions", 1)
    add_numbered(
        doc,
        [
            "Create a one-page supervisor proposal using the title, RQs, contributions, and minimum viable scope.",
            "Finalize the public JSONL trace schema and keep it stable before running full experiments.",
            "Implement the custom ReAct pilot first, because it validates the schema without API or framework risk.",
            "Add LangChain instrumentation after the custom pilot works.",
            "Run a small pilot and inspect whether retrieval, memory, and tool traces are labeled correctly.",
            "Only then scale to the full experiment matrix and thesis figures.",
        ],
    )

    add_heading(doc, "7. Risk Controls", 1)
    add_table(
        doc,
        ["Risk", "Default Control", "Fallback"],
        [
            ["Second framework takes too long", "Use custom ReAct as the second implementation.", "Drop LlamaIndex or any heavy second framework."],
            ["Provider API is unstable or expensive", "Use fixed tasks and token/cost proxy.", "Reduce repetitions and keep mock/custom traces for method validation."],
            ["Mitigation effect is weak", "Frame guard as detection-first.", "Make Chapter 7 a detection-only case study."],
            ["Scope becomes too broad", "Prioritize schema, toolkit, and empirical analysis.", "Remove dashboard, complex success scoring, and duplicate suppression."],
        ],
        [2500, 3430, 3430],
    )

    add_heading(doc, "8. Definition of Done", 1)
    add_bullets(
        doc,
        [
            "The artifact can capture and export model-visible context for each invocation.",
            "The provenance schema distinguishes the main context sources needed by the thesis.",
            "The analysis produces interpretable tables and figures for all research questions.",
            "The guard produces meaningful flags for growth, duplication, and source dominance.",
            "The thesis clearly shows how the seminar paper evolved from observation into system, method, and empirical evaluation.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(footer, "MSc Thesis Action Plan", color=MUTED, size=9)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print("wrote action plan docx")
