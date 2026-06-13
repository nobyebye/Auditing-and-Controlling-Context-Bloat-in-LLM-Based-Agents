from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from scripts.build_action_plan_docx import (
    DARK_BLUE,
    INK,
    MUTED,
    add_bullets,
    add_callout,
    add_heading,
    add_numbered,
    add_table,
    add_text,
    configure_document,
)


OUT = Path("Context_Bloat_Action_Plan.docx")


def build() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(title, "MSc Thesis Action Plan", bold=True, color=INK, size=22)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    add_text(subtitle, "Auditing and Controlling Context Bloat in LLM-Based Agents", color=MUTED, size=12)

    meta = doc.add_paragraph()
    add_text(meta, "Chinese title: ", bold=True, color=DARK_BLUE)
    add_text(meta, "面向 LLM 智能体上下文膨胀的检测、量化与缓解研究")

    add_callout(
        doc,
        "Core shift:",
        "Automatically constructed context remains the background and object of observation, but context bloat is the central research problem. Runtime tracing and context auditing become the method for studying bloat.",
    )

    add_heading(doc, "1. Thesis Focus", 1)
    add_bullets(
        doc,
        [
            "Study context bloat in LLM-based agents rather than general context auditing alone.",
            "Detect and localize bloated context inside retrieval, memory, tool traces, conversation history, and framework-generated content.",
            "Measure redundancy, growth, uniqueness, and source contribution using reproducible metrics.",
            "Evaluate whether simple mitigation methods reduce unnecessary context without significantly harming task performance.",
            "Preserve the seminar paper foundation: runtime tracing, provenance schema, and model-visible context analysis.",
        ],
    )

    add_heading(doc, "2. Research Questions", 1)
    add_numbered(
        doc,
        [
            "How can context bloat be detected and localized in LLM-based agents?",
            "How can context bloat be quantitatively measured?",
            "What are the main sources and patterns of context bloat?",
            "Can context bloat be mitigated without significantly reducing task performance?",
        ],
    )

    add_heading(doc, "3. Work Packages", 1)
    add_table(
        doc,
        ["WP", "Work Package", "Concrete Output", "Acceptance Check"],
        [
            ["1", "Bloat definition and taxonomy", "Definition of context bloat plus categories such as duplicate retrieval, stale memory, repeated tool traces, irrelevant history, oversized passages, and framework overhead.", "The thesis has a problem taxonomy, not only an auditing framework description."],
            ["2", "Runtime tracing foundation", "Trace schema, capture wrapper, JSONL exporter, and provenance labels for model-visible messages.", "Every invocation is captured and each segment has a source label."],
            ["3", "Bloat metrics", "Redundancy Ratio, Unique Information Ratio, Context Growth Rate, Source Contribution Ratio, duplicate count, dominance flags, and cost proxy.", "Metrics can be computed automatically from JSONL traces."],
            ["4", "Localization analysis", "Per-source bloat reports for retrieval, memory, tool traces, conversation history, and framework content.", "The analysis can answer where bloat comes from in each workflow."],
            ["5", "Controlled workflows", "Retrieval QA, multi-step tool use, memory turns, and combined retrieval-memory-tool tasks.", "Each workflow creates traces suitable for detecting bloat patterns."],
            ["6", "Mitigation case study", "Exact duplicate removal plus one optional method: memory filtering or tool-output compression.", "Reduced traces show token/redundancy changes and task performance impact."],
            ["7", "Empirical results", "Tables and figures showing bloat patterns, sources, growth rates, and mitigation trade-offs.", "Results directly answer RQ1-RQ4."],
            ["8", "Writing and final artifact", "Chapters 1-9, artifact README, reproducible scripts, and final figures.", "The thesis presents context bloat as the main problem and auditing as the method."],
        ],
        [850, 2200, 3150, 3160],
    )

    add_heading(doc, "4. Metrics to Implement", 1)
    add_table(
        doc,
        ["Metric", "Meaning", "Use in Thesis"],
        [
            ["Redundancy Ratio", "Share of context tokens belonging to repeated or redundant segments.", "Measures how much context is likely unnecessary."],
            ["Unique Information Ratio", "Share of context tokens that remain after redundant content is collapsed.", "Shows whether context growth adds new information or repeats old content."],
            ["Context Growth Rate", "Relative token increase across successive LLM invocations.", "Detects rapid accumulation during multi-step workflows."],
            ["Source Contribution Ratio", "Share of total or redundant context contributed by each source type.", "Localizes bloat to memory, retrieval, tools, conversation history, or framework content."],
            ["Task Performance Check", "Lightweight correctness/usefulness judgment before and after mitigation.", "Tests whether mitigation damages utility."],
        ],
        [2100, 3800, 3460],
    )

    add_heading(doc, "5. 18-Week Execution Timeline", 1)
    add_table(
        doc,
        ["Weeks", "Focus", "Actions", "Deliverables"],
        [
            ["1-2", "Reframe proposal", "Rewrite title, RQs, contributions, and chapter outline around context bloat.", "Supervisor-ready bloat-focused proposal."],
            ["3-4", "Taxonomy and schema", "Define bloat types and update provenance schema to support localization.", "Conceptual model and trace schema."],
            ["5-6", "Metrics implementation", "Add redundancy, uniqueness, growth, source contribution, and cost metrics.", "Metric scripts and unit tests."],
            ["7-8", "Controlled workflows", "Prepare retrieval, memory, tool, conversation-history, and combined tasks.", "Pilot task set and validated traces."],
            ["9-11", "Main experiments", "Run configurations, collect traces, summarize bloat sources and patterns.", "Dataset, tables, and figures for RQ1-RQ3."],
            ["12-13", "Mitigation study", "Test duplicate removal and one optional mitigation method.", "Before/after mitigation comparison for RQ4."],
            ["14-16", "Writing", "Draft chapters on model, method, metrics, setup, results, and mitigation.", "Complete technical draft."],
            ["17-18", "Revision and polish", "Strengthen discussion, limitations, references, artifact README, and slides.", "Submission-ready thesis package."],
        ],
        [1050, 1750, 4050, 2510],
    )

    add_heading(doc, "6. Immediate Next Actions", 1)
    add_numbered(
        doc,
        [
            "Rename the thesis direction to context bloat and update the one-page proposal.",
            "Move runtime tracing, provenance schema, and auditing into the Method section rather than presenting them as the final goal.",
            "Define the bloat taxonomy before running large experiments.",
            "Implement or verify the four core metrics: Redundancy Ratio, Unique Information Ratio, Context Growth Rate, and Source Contribution Ratio.",
            "Run a small pilot to confirm the tool can localize bloat by source type.",
            "Choose one conservative mitigation method for the main evaluation: exact duplicate removal is the safest default.",
        ],
    )

    add_heading(doc, "7. Risk Controls", 1)
    add_table(
        doc,
        ["Risk", "Default Control", "Fallback"],
        [
            ["Bloat definition becomes vague", "Use a concrete taxonomy with detectable patterns.", "Limit claims to repeated, redundant, stale, irrelevant, or oversized context."],
            ["Mitigation hurts task quality", "Evaluate mitigation conservatively and report trade-offs.", "Make mitigation a small case study rather than a broad optimization claim."],
            ["Near-duplicate detection is hard", "Start with normalized exact hashes.", "Discuss semantic redundancy as future work."],
            ["Scope expands too much", "Keep runtime tracing as method and context bloat as the problem.", "Drop dashboard, large benchmark, and multiple mitigation methods."],
        ],
        [2500, 3430, 3430],
    )

    add_heading(doc, "8. Definition of Done", 1)
    add_bullets(
        doc,
        [
            "The thesis clearly states context bloat as the central problem.",
            "The artifact detects and localizes bloat by source type.",
            "The metrics quantify redundancy, uniqueness, growth, and source contribution.",
            "The experiments identify which workflows and sources create the most bloat.",
            "The mitigation case study reports both context reduction and task performance impact.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(footer, "Context Bloat Thesis Action Plan", color=MUTED, size=9)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print("wrote Context_Bloat_Action_Plan.docx")
